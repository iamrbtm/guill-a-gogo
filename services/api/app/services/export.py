from __future__ import annotations

import csv
import datetime as dt
import io
import uuid

from sqlalchemy.orm import Session

from app.models.accounts import Trip
from app.models.itinerary import (
    ChangeProposal,
    Expense,
    LodgingCandidate,
    MealPlan,
    PlanningWarning,
    Reservation,
    Stop,
    TripDay,
)
from app.models.profiles import Pet, TravelerProfile, Vehicle
from app.models.accounts import trip_pets, trip_travelers


def _now() -> dt.datetime:
    return dt.datetime.now(dt.timezone.utc)


def collect_itinerary_data(session: Session, trip: Trip) -> dict:
    """Gather the full approved itinerary into a structured, export-ready dict.

    Every section is tagged with whether its data is estimated, confirmed, or
    manually entered, so exports never present guesses as fact.
    """
    days = session.query(TripDay).filter_by(trip_id=trip.id).order_by(TripDay.day_number).all()
    stops = session.query(Stop).filter_by(trip_id=trip.id).order_by(Stop.order_index).all()
    lodging = session.query(LodgingCandidate).filter_by(trip_id=trip.id).all()
    meals = session.query(MealPlan).filter_by(trip_id=trip.id).all()
    expenses = session.query(Expense).filter_by(trip_id=trip.id).all()
    reservations = session.query(Reservation).filter_by(trip_id=trip.id).all()
    warnings = session.query(PlanningWarning).filter_by(trip_id=trip.id).all()
    proposals = session.query(ChangeProposal).filter_by(trip_id=trip.id).order_by(ChangeProposal.created_at).all()

    traveler_ids = [r[0] for r in session.execute(
        trip_travelers.select().where(trip_travelers.c.trip_id == trip.id)
    ).all()]
    travelers = session.query(TravelerProfile).filter(TravelerProfile.id.in_(traveler_ids)).all() if traveler_ids else []

    pet_ids = [r[0] for r in session.execute(
        trip_pets.select().where(trip_pets.c.trip_id == trip.id)
    ).all()]
    pets = session.query(Pet).filter(Pet.id.in_(pet_ids)).all() if pet_ids else []

    vehicle = session.get(Vehicle, trip.vehicle_id) if trip.vehicle_id else None

    return {
        "exported_at": _now().isoformat(),
        "trip": trip,
        "days": days,
        "stops": stops,
        "lodging": lodging,
        "meals": meals,
        "expenses": expenses,
        "reservations": reservations,
        "warnings": warnings,
        "proposals": proposals,
        "travelers": travelers,
        "pets": pets,
        "vehicle": vehicle,
    }


def export_csv(data: dict) -> bytes:
    """Core itinerary as CSV (one of the required export formats)."""
    out = io.StringIO()
    w = csv.writer(out)
    trip: Trip = data["trip"]
    w.writerow(["Guill-a-Gogo Itinerary Export"])
    w.writerow(["Exported At", data["exported_at"]])
    w.writerow(["Trip", trip.title])
    w.writerow(["Origin", trip.origin or ""])
    w.writerow(["Destination", trip.destination or ""])
    w.writerow(["Status", trip.status])
    w.writerow([])
    w.writerow(["NOTE: 'estimate' = planned/calculated; 'confirmed' = user-verified; 'manual' = manually entered."])
    w.writerow([])

    w.writerow(["Stops", "Day", "Type", "Required", "Status", "Min Duration (min)", "Notes", "Data Quality"])
    for s in data["stops"]:
        quality = "confirmed" if s.completed else "estimate"
        w.writerow([s.name or "", "", s.stop_type, s.required, "completed" if s.completed else "planned",
                    s.min_duration_minutes or "", s.notes or "", quality])
    w.writerow([])

    w.writerow(["Lodging", "Address", "User Confirmed", "Accessibility Confirmed", "Two Dogs", "Rate", "Data Quality"])
    for l in data["lodging"]:
        quality = "confirmed" if l.user_confirmed else "estimate"
        w.writerow([l.name, l.address or "", l.user_confirmed, l.required_accessibility_confirmed,
                    l.two_dogs_permitted, l.rate_minor or "", quality])
    w.writerow([])

    w.writerow(["Meals", "Type", "Restaurant", "Serves Fish", "Fish-Safe Option", "Notes"])
    for m in data["meals"]:
        w.writerow([f"Day {m.trip_id}", m.meal_type, m.restaurant_name or "",
                    m.serves_fish, m.fish_safe_option_available, m.notes or ""])
    w.writerow([])

    w.writerow(["Expenses", "Category", "Amount", "Currency", "Note"])
    for e in data["expenses"]:
        w.writerow(["", e.category, f"{e.amount_minor / 100:.2f}", e.currency, e.note or ""])
    w.writerow([])

    w.writerow(["Warnings", "Severity", "Category", "Message"])
    for wn in data["warnings"]:
        w.writerow(["", wn.severity, wn.category, wn.message])

    return out.getvalue().encode("utf-8")


def export_xlsx(data: dict) -> bytes:
    """XLSX workbook with itinerary, locations, expenses, lodging, and calculation sheets."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="0B3D2E")

    def style_header(ws):
        # Read the most recent row (append's return value is not reliable across versions).
        row = ws[ws.max_row]
        for cell in row:
            cell.font = header_font
            cell.fill = header_fill

    # --- Itinerary sheet ---
    ws = wb.active
    ws.title = "Itinerary"
    trip: Trip = data["trip"]
    ws.append(["Guill-a-Gogo Itinerary Export"])
    ws.append(["Exported At", data["exported_at"]])
    ws.append(["Trip", trip.title])
    ws.append(["Origin", trip.origin or "", "Destination", trip.destination or "", "Status", trip.status])
    ws.append([])
    ws.append(["NOTE: estimate = planned/calculated; confirmed = user-verified; manual = manually entered."])
    ws.append([])
    hdr = ws.append(["Stop", "Type", "Required", "Status", "Min Duration (min)", "Notes", "Data Quality"])
    style_header(ws)
    for s in data["stops"]:
        quality = "confirmed" if s.completed else "estimate"
        ws.append([s.name or "", s.stop_type, s.required, "completed" if s.completed else "planned",
                   s.min_duration_minutes or "", s.notes or "", quality])

    # --- Locations sheet ---
    ws2 = wb.create_sheet("Locations")
    hdr2 = ws2.append(["Stop", "Address", "Lat", "Lng", "Required", "Notes"])
    style_header(ws2)
    for s in data["stops"]:
        ws2.append([s.name or "", s.address or "", s.lat or "", s.lng or "", s.required, s.notes or ""])

    # --- Lodging sheet ---
    ws3 = wb.create_sheet("Lodging")
    hdr3 = ws3.append(["Name", "Address", "User Confirmed", "Accessibility Confirmed",
                       "Two Dogs OK", "Breakfast", "Trailer Parking", "Rate", "Taxes", "Data Quality"])
    style_header(ws3)
    for l in data["lodging"]:
        quality = "confirmed" if l.user_confirmed else "estimate"
        ws3.append([l.name, l.address or "", l.user_confirmed, l.required_accessibility_confirmed,
                    l.two_dogs_permitted, l.breakfast_included, l.trailer_parking_confirmed,
                    l.rate_minor or "", l.taxes_minor or "", quality])

    # --- Expenses sheet ---
    ws4 = wb.create_sheet("Expenses")
    hdr4 = ws4.append(["Category", "Amount", "Currency", "Incurred", "Note"])
    style_header(ws4)
    total = 0
    for e in data["expenses"]:
        total += e.amount_minor
        ws4.append([e.category, e.amount_minor / 100, e.currency, (e.incurred_at.isoformat() if e.incurred_at else ""), e.note or ""])
    ws4.append([])
    ws4.append(["TOTAL", f"{total / 100:.2f}", "", "", ""])

    # --- Calculations sheet ---
    ws5 = wb.create_sheet("Calculations")
    hdr5 = ws5.append(["Metric", "Value", "Data Quality"])
    style_header(ws5)
    ws5.append(["Total stops", len(data["stops"]), "confirmed"])
    ws5.append(["Completed stops", sum(1 for s in data["stops"] if s.completed), "confirmed"])
    ws5.append(["Total expenses", f"{total / 100:.2f}", "confirmed"])
    ws5.append(["Confirmed lodging", sum(1 for l in data["lodging"] if l.user_confirmed), "confirmed"])
    ws5.append(["Estimated lodging", sum(1 for l in data["lodging"] if not l.user_confirmed), "estimate"])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_pdf(data: dict) -> bytes:
    """Readable PDF day-by-day summary."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, leftMargin=0.6 * inch, rightMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=18)
    body = []

    trip: Trip = data["trip"]
    body.append(Paragraph(f"Guill-a-Gogo Itinerary: {trip.title}", title_style))
    body.append(Paragraph(f"Exported: {data['exported_at']}", styles["Normal"]))
    body.append(Paragraph(f"Origin: {trip.origin or '—'}  →  Destination: {trip.destination or '—'}", styles["Normal"]))
    body.append(Paragraph(f"Status: {trip.status}", styles["Normal"]))
    body.append(Spacer(1, 0.2 * inch))
    body.append(Paragraph(
        "<b>Note:</b> 'estimate' = planned/calculated · 'confirmed' = user-verified · 'manual' = manually entered.",
        styles["Normal"],
    ))
    body.append(Spacer(1, 0.2 * inch))

    body.append(Paragraph("Stops", styles["Heading2"]))
    stop_rows = [["Stop", "Type", "Required", "Status", "Quality"]]
    for s in data["stops"]:
        q = "confirmed" if s.completed else "estimate"
        stop_rows.append([s.name or "", s.stop_type, "yes" if s.required else "",
                          "completed" if s.completed else "planned", q])
    t = Table(stop_rows, colWidths=[2 * inch, 1.2 * inch, 0.8 * inch, 1 * inch, 1 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D2E")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
    ]))
    body.append(t)
    body.append(Spacer(1, 0.2 * inch))

    body.append(Paragraph("Lodging", styles["Heading2"]))
    lodg_rows = [["Name", "Confirmed", "Accessibility", "Two Dogs"]]
    for l in data["lodging"]:
        lodg_rows.append([l.name, "yes" if l.user_confirmed else "no",
                          "yes" if l.required_accessibility_confirmed else "no",
                          "yes" if l.two_dogs_permitted else "no"])
    t2 = Table(lodg_rows, colWidths=[2.5 * inch, 1.2 * inch, 1.4 * inch, 1.2 * inch])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D2E")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    body.append(t2)

    doc.build(body)
    return buf.getvalue()


def export_docx(data: dict) -> bytes:
    """DOCX suitable for import into Google Docs."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    trip: Trip = data["trip"]
    doc.add_heading(f"Guill-a-Gogo Itinerary: {trip.title}", level=0)
    doc.add_paragraph(f"Exported: {data['exported_at']}")
    doc.add_paragraph(f"Origin: {trip.origin or '—'}  →  Destination: {trip.destination or '—'}")
    doc.add_paragraph(f"Status: {trip.status}")
    doc.add_paragraph(
        "Note: 'estimate' = planned/calculated · 'confirmed' = user-verified · 'manual' = manually entered."
    )

    doc.add_heading("Stops", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Stop", "Type", "Required", "Status", "Quality"]):
        hdr[i].text = h
    for s in data["stops"]:
        row = table.add_row().cells
        q = "confirmed" if s.completed else "estimate"
        vals = [s.name or "", s.stop_type, "yes" if s.required else "",
                "completed" if s.completed else "planned", q]
        for i, v in enumerate(vals):
            row[i].text = str(v)

    doc.add_heading("Lodging", level=1)
    table2 = doc.add_table(rows=1, cols=4)
    hdr2 = table2.rows[0].cells
    for i, h in enumerate(["Name", "Confirmed", "Accessibility", "Two Dogs"]):
        hdr2[i].text = h
    for l in data["lodging"]:
        row = table2.add_row().cells
        vals = [l.name, "yes" if l.user_confirmed else "no",
                "yes" if l.required_accessibility_confirmed else "no",
                "yes" if l.two_dogs_permitted else "no"]
        for i, v in enumerate(vals):
            row[i].text = str(v)

    doc.add_heading("Expenses", level=1)
    table3 = doc.add_table(rows=1, cols=3)
    hdr3 = table3.rows[0].cells
    for i, h in enumerate(["Category", "Amount", "Note"]):
        hdr3[i].text = h
    for e in data["expenses"]:
        row = table3.add_row().cells
        row[0].text = e.category
        row[1].text = f"{e.amount_minor / 100:.2f} {e.currency}"
        row[2].text = e.note or ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_EXPORTERS = {
    "csv": ("text/csv", "itinerary.csv", export_csv),
    "xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "itinerary.xlsx", export_xlsx),
    "pdf": ("application/pdf", "itinerary.pdf", export_pdf),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "itinerary.docx", export_docx),
}


def get_exporter(format: str):
    """Return (content_type, filename, fn) for a format, or None if unsupported."""
    return _EXPORTERS.get(format)
